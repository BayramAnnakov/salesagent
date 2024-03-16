from dotenv import load_dotenv

load_dotenv()


import logging
from telegram import Update
from telegram.ext import filters
from telegram.ext import ApplicationBuilder, ContextTypes, CommandHandler, MessageHandler

import os
import time

from llama_index.agent.openai import OpenAIAgent

from agent import get_openai_agent

from docx import Document

from circle import create_transfer



logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)

agent = get_openai_agent()

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await context.bot.send_message(chat_id=update.effective_chat.id, text="Checking the upcoming sales calendar events on March 17th 2024... 📅")
    sticker_message = await context.bot.send_message(chat_id=update.effective_chat.id, text="🤖")
    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="typing")

    response = agent.chat("Search the upcoming sales calendar events only on March 17th 2024. Format the response for Telegram message, use emoji.")
    await context.bot.delete_message(chat_id=update.effective_chat.id, message_id=sticker_message.message_id)
    await context.bot.send_message(chat_id=update.effective_chat.id, text=str(response), parse_mode="Markdown")

    response = agent.chat("Create onchain sales job.")
    await context.bot.send_message(chat_id=update.effective_chat.id, text=str(response), parse_mode="Markdown")

    await context.bot.send_message(chat_id=update.effective_chat.id, text="Fetching LinkedIn profile of the event participant and preparing memo... 🕵️‍♂️")
    
    response = agent.chat("Prepare a memo how to prepare for this private jet services sales call using info about the event participant from their LinkedIn profile. Score this lead's success probability from 1 to 10 based on the LinkedIn profile information and the upcoming sales call event details. List possible topics or questions to discuss/ask to make the sales call successful. Format the response for Telegram message, use emoji.")
    await context.bot.send_message(chat_id=update.effective_chat.id, text=str(response), parse_mode="Markdown")
    
    await context.bot.send_message(chat_id=update.effective_chat.id, text="Transferring USDC to the LinkedIn Agent for LinkedIn profile fetch... 💸")
    transaction_hash = create_transfer("0.1", "0xeb6e084738dff0739655a99df0de4f37ce979a71", "f1c83e00-19b9-5feb-adfe-0e3de5ebaf29")
    await context.bot.send_message(chat_id=update.effective_chat.id, text="Transfer complete, transaction ID: "+transaction_hash+" 🎉")


    #sleep for 10 seconds
    await context.bot.send_message(chat_id=update.effective_chat.id, text="😴")
    time.sleep(10)

    await context.bot.send_message(chat_id=update.effective_chat.id, text="Detected meeting transcript, evaluating... 🕵️‍♂️")  

    response = agent.chat("Analyze the sales call using meeting transcript that is downloaded by zoom meeting id. Score the sales call from 1 to 10. Format response for Telegram message, use emoji")
    await context.bot.send_message(chat_id=update.effective_chat.id, text=str(response), parse_mode="Markdown")
    # meeting_analysis_document = str(response)
    
    # doc = Document()

    # doc.add_heading('Meeting Analysis', 0)
    # doc.add_paragraph(meeting_analysis_document)
    # doc.save('meeting_analysis.docx')

    # with open('meeting_analysis.docx', 'rb') as file:
    #     await context.bot.send_document(chat_id=update.effective_chat.id, document=file, caption="📋 I've analyzed the sales call. Here's the meeting analysis document.")

    response = agent.chat("Complete the onchain sales job with the sales call performance score. Format the response for Telegram message, use emoji.")
    await context.bot.send_message(chat_id=update.effective_chat.id, text=str(response), parse_mode="Markdown")
    
    await context.bot.send_message(chat_id=update.effective_chat.id, text="Transferring USDC to the AI sales agent for the successful evaluation of the sales call... 💸")
    transaction_hash = create_transfer("0.5", "0xd1c31e2c6c5558c306c9c71d51e1faffd80ef517", "0b28bb7d-7584-5585-ad56-12a3c814d427")
    await context.bot.send_message(chat_id=update.effective_chat.id, text="Transfer complete, transaction ID: "+transaction_hash+" 🎉")

    response = agent.chat("Update the CRM with the sales call score and topics discussed. Format the response for Telegram message, use emoji.")
    await context.bot.send_message(chat_id=update.effective_chat.id, text=str(response), parse_mode="Markdown")

if __name__ == '__main__':
    application = ApplicationBuilder().token(os.environ["TELEGRAM_BOT_TOKEN"]).build()

    start_handler = CommandHandler('start', start)
    application.add_handler(start_handler)
    
    application.run_polling()